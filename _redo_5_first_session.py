"""
§5 The First Session cleanup:

1. Delete the @Claude note + the empty List Bullet rows + the orphaned
   'December' / '2023' artifact.
2. Drop the floating 'Initial Session: (2 hours-Bill IDA/BPS, CFARS (6+),
   TP or therapy session)' line - the codes/CFARS portion duplicates
   §3.3, but we keep the '2 hours' scheduling piece by merging it into
   the Medicaid/private-insurance timing bullet (preface 'Block 2 hours.').
3. Convert the two floating List Paragraph items ('Talk with caregiver...'
   and 'If there is sensitive information...') into proper List Bullets
   and move them to sit between 'child must be present' and 'using as
   IDA/BPS, see §6'.
4. Insert a new Heading 2 'How the First Session Runs' right after the §5
   heading so the operational bullets have a label, parallel to the
   existing 'Topics to Cover in the First Session' subheading below.
"""
from docx import Document
from copy import deepcopy

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)

# Anchors / targets
heading_5_el = None
claude_note_el = None
initial_session_el = None  # ¶138 floating (to delete, after merging timing)
talk_caregiver_p = None    # ¶139 (move + restyle)
sensitive_info_p = None    # ¶140 (move + restyle)
december_el = None         # ¶141
twenty_three_el = None     # ¶142
medicaid_timing_p = None   # ¶145 (modify)
child_present_p = None     # ¶146 (anchor for moving 139/140)
ida_bps_see_s6_p = None    # ¶147 (sanity-check anchor)
empties_in_section_5 = []  # all empty list-bullet/paragraph rows in §5

in_s5 = False
for p in doc.paragraphs:
    t = p.text
    if "5. The First Session" in t and p.style.name.startswith("Heading"):
        heading_5_el = p._element
        in_s5 = True
        continue
    if in_s5 and t.startswith("Topics to Cover in the First Session"):
        in_s5 = False
        break
    if not in_s5:
        continue

    stripped = t.strip()
    if claude_note_el is None and "@Claude, add this" in t:
        claude_note_el = p._element
        continue
    if initial_session_el is None and "Initial Session: (2 hours" in t:
        initial_session_el = p._element
        continue
    if talk_caregiver_p is None and "Talk with caregiver about presenting problems" in t:
        talk_caregiver_p = p
        continue
    if sensitive_info_p is None and "If there is sensitive information that the caregiver does not want to discuss" in t:
        sensitive_info_p = p
        continue
    if december_el is None and stripped == "December":
        december_el = p._element
        continue
    if twenty_three_el is None and stripped == "2023":
        twenty_three_el = p._element
        continue
    if medicaid_timing_p is None and "For Medicaid, actual time with child and family" in t:
        medicaid_timing_p = p
        continue
    if child_present_p is None and "The child must be present for at least part of the session" in t:
        child_present_p = p
        continue
    if ida_bps_see_s6_p is None and "If you are using the first session as an IDA or BPS" in t:
        ida_bps_see_s6_p = p
        continue
    if stripped == "" and (p.style.name.startswith("List") or p.style.name == "Normal"):
        empties_in_section_5.append(p._element)

# Sanity checks
problems = []
for name, val in [
    ("§5 heading", heading_5_el),
    ("@Claude note", claude_note_el),
    ("Initial Session line", initial_session_el),
    ("Talk-with-caregiver line", talk_caregiver_p),
    ("Sensitive-info line", sensitive_info_p),
    ("'December' artifact", december_el),
    ("'2023' artifact", twenty_three_el),
    ("Medicaid timing bullet", medicaid_timing_p),
    ("Child-must-be-present bullet", child_present_p),
    ("IDA/BPS-see-§6 bullet", ida_bps_see_s6_p),
]:
    if val is None:
        problems.append(name)
if problems:
    raise SystemExit(f"Could not locate: {problems}")
print(f"Found §5 anchors. {len(empties_in_section_5)} empty filler rows to delete.")

# ---------- 1. Modify the Medicaid timing bullet to absorb '2 hours' ----------
new_timing = ("Block 2 hours. For Medicaid, actual time with child and "
              "family may be 1.5 hours; for private insurance, 60 minutes.")
for r in list(medicaid_timing_p.runs):
    r._element.getparent().remove(r._element)
medicaid_timing_p.add_run(new_timing)
print("Updated Medicaid timing bullet to: " + new_timing)

# ---------- 2. Move 'Talk with caregiver' + 'Sensitive info' into place ----------
# Detach both, re-style as List Bullet, insert after child_present_p in order.
def detach(p):
    el = p._element
    el.getparent().remove(el)
    return el

talk_el = detach(talk_caregiver_p)
sensitive_el = detach(sensitive_info_p)

# Set style to "List Bullet"
talk_caregiver_p.style = doc.styles["List Bullet"]
sensitive_info_p.style = doc.styles["List Bullet"]

# Insert: child_present -> talk -> sensitive (then ida_bps_see_s6 follows)
child_present_p._element.addnext(sensitive_el)  # this places sensitive right after child_present
sensitive_el.addprevious(talk_el)               # then talk goes between child_present and sensitive
# Now the order is: child_present_p → talk_el → sensitive_el → ida_bps_see_s6_p

# ---------- 3. Insert new Heading 2 'How the First Session Runs' ----------
how_h = doc.add_heading("How the First Session Runs", level=2)
heading_5_el.addnext(how_h._element)

# ---------- 4. Delete noise: @Claude, Initial Session line, December, 2023, empties ----------
to_delete = (
    [claude_note_el, initial_session_el, december_el, twenty_three_el]
    + empties_in_section_5
)
removed = 0
for el in to_delete:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)
        removed += 1

doc.save(PATH)
print(f"§5 done: removed {removed} noise/empty paragraphs; "
      f"moved 2 floating bullets into place; added 'How the First Session Runs' subheading.")
