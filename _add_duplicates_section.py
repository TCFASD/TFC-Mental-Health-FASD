"""
Append a 'Duplicated / Overlapping Content' section to the end of the
REVISED manual. Every line of duplicated content gets a yellow highlight.
Group headers are bold and not highlighted. Originals are left in place.
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_normal(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_highlighted(text, style=None):
    """Add a paragraph with the given text highlighted yellow."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def add_group(title, intro, items):
    """
    title: e.g. "1. CFARS rules"
    intro: short text explaining the conflict / overlap (not highlighted)
    items: list of (location_label, text) tuples - text is highlighted
    """
    add_heading(title, level=2)
    if intro:
        add_normal(intro, italic=True)
    for loc, text in items:
        # location label is bold + not highlighted, then highlighted text
        p = doc.add_paragraph()
        loc_run = p.add_run(f"[{loc}] ")
        loc_run.bold = True
        body = p.add_run(text)
        body.font.highlight_color = WD_COLOR_INDEX.YELLOW


# ----- Section header -----
doc.add_page_break()
add_heading("Duplicated / Overlapping Content - For You to Review", level=1)
add_normal(
    "This appendix collects every place in the manual where the same topic "
    "is covered more than once. Each group below lists every version of the "
    "content that currently exists in the document, with the location shown "
    "in bold. The duplicated text itself is highlighted yellow. The originals "
    "have NOT been removed - decide which version you want to keep, and the "
    "others can be deleted from their in-section locations.",
    italic=True,
)
add_normal(
    "Note: groups marked CONFLICT contain rules that disagree with each "
    "other, not just repeat each other. Those need a decision, not just a "
    "dedup.",
    italic=True,
)

# ===== 1. CFARS =====
add_group(
    "1. CFARS rules",
    "Appears in 4 places. The §11 version is nearly verbatim with §6. The newer billing workflow uses ages 5-17 (with FARS for 18+); the older sections say 'over 6 years of age.'",
    [
        ("§3.2 Service Codes",
         "H0031 (no modifier) - CFARS (Limited Functional Assessment). Medicaid only. The first 3 per state fiscal year are billed here; the 4th is billed to CCM. Client must be present."),
        ("§6 Required Elements at Assessment",
         "CFARS must be completed. CFARS is due at initiation of services and quarterly for Medicaid clients over 6 years of age (up to 3 times per year). Interns are not able to bill for CFARS. CFARS are completed in Lauris."),
        ("§11 Procedures Summary - CFARS",
         "CFARS are due at initiation of services and quarterly for Medicaid clients over 6 years of age, up to 3 times a year. Interns cannot bill for CFARS. CFARS are completed in Lauris. Reference: CFARS (Florida DCF)."),
        ("§FASD Billing Workflow - CFARS",
         "Required for clients ages 5-17 (use FARS for adults 18+). Administered at clinical discretion; no separate annual billing cap. State outcome data is reported semi-annually (January 15 and July 15). Must be administered by an individual certified by DCF to administer the instrument."),
    ],
)

# ===== 2. PCP letter =====
add_group(
    "2. PCP / Physician letter",
    "Appears in 3 places. §11 adds a 'Responsibility' line; otherwise these say the same thing.",
    [
        ("§3.9 What Intake Does",
         "All primary care providers should receive notice of services. Ensure the release to communicate with the PCP is signed."),
        ("§4 Consent Forms",
         "Have the family sign the Primary Care Physician (PCP) release so the PCP can be notified of services."),
        ("§11 Physician Letter",
         "All primary care providers should receive notice of services. Ensure the release to communicate with the PCP is signed. Responsibility: Therapists and Interns."),
    ],
)

# ===== 3. Outcome measures =====
add_group(
    "3. Outcome measures",
    "Appears in 5 places. The lists of which measures to use don't match exactly: §5 says ACEs + PSS + PSOC; §6 adds the Neurobehavioral Screener; §9 says PSOC + PSS + ACE short form. Pick one canonical list.",
    [
        ("§5 First Session topics",
         "Fill out outcome measures (see Section 9). If the family came through the clinic you may already have most of these. If not, include ACEs, Perceived Stress Scale and the Parenting Sense of Competency tool."),
        ("§6 Assessments - Timelines",
         "Outcome measures: ACE score, Perceived Stress Scale, Parenting Sense of Competency (in Lauris), and Neurobehavioral Screener for both caregivers and child."),
        ("§6 Outcome Measures Location",
         "All outcome measures are stored in Box under MH Therapist / Outcome Measures: Outcome Measures (Box)."),
        ("§9 Outcomes",
         "Outcomes must be completed at the beginning of services to establish an accurate baseline. When you get the case, put the outcome-measure dates on your calendar. They are required at initiation, every 3 months, and at closing. Outcomes to complete: Parenting Sense of Competency, Perceived Stress Scale, and ACE short form (only at the beginning of treatment)."),
        ("Appendix B Outcome Measures",
         "MH Therapist / Outcome Measures (Box)"),
    ],
)

# ===== 4. TP deadline (CONFLICT) =====
add_group(
    "4. Treatment-plan deadline - CONFLICT (10 days vs. 45 days)",
    "These two rules disagree. Decide which one is current and delete the other.",
    [
        ("§3 Billing Absolute Rule #3",
         "Rule #3. You must complete a Treatment Plan within 10 calendar days from the date of the assessment, and the caregiver must sign it. (Per Rule 65D-30.0044, F.A.C. Note: this is calendar days, not business days.)"),
        ("§7 Treatment Plan Deadline in Practice",
         "The Treatment Plan must be completed and reviewed with the caregiver within the 10-calendar-day window. The caregiver's signature does NOT have to be received within the 10-calendar-day window."),
        ("§FASD Billing Workflow - Documentation Checklist",
         "Treatment plan signed by the treating practitioner within 45 days of the start of services (Medicaid reimburses for services provided within 45 days prior to the practitioner signature)."),
    ],
)

# ===== 5. BSE same-day rule + BSE caps =====
add_group(
    "5. BSE / BBSE rules",
    "Appears in 2-3 places. §FASD Billing Workflow has the most detail; §3.2 and §6 have the basics. Decide whether the detailed version belongs in §3 or §6.",
    [
        ("§3.2 H2010 HO",
         "H2010 HO - BBSE (Brief Behavioral Status Exam). Licensed clinicians only. 1-2 units (15 or 30 minutes). Maximum 2 units per day, maximum 10 units per year. Cannot be billed the same day as PSY or IDA."),
        ("§6 Required Elements",
         "A Brief Status Exam (BSE) is required. Per the Oct 2024 Summary of Procedures: the BSE must be completed by a licensed practitioner prior to the Treatment Plan when the assessment is done by an unlicensed therapist. If the assessment is done by a licensed person, the BSE may be incorporated into the assessment. The BSE cannot be billed the same day as the assessment."),
        ("§FASD Billing Workflow - Decision Rule",
         "If the FASD evaluation is older than 6 months, or was not signed by a qualifying practitioner, complete a Brief Behavioral Health Status Examination (BSE, H2010 HO) before treatment plan development. Never bill a BSE on the same day as a psychiatric evaluation, biopsychosocial evaluation, or in-depth assessment for the same recipient."),
        ("§FASD Billing Workflow - BSE Constraints",
         "Code: H2010 HO. Telemedicine: H2010 HO GT. Maximum 2 quarter-hour units (30 minutes) per day. Maximum 10 quarter-hour units (2.5 hours) per recipient per state fiscal year. Required components: purpose of the exam, mental health status, summary of findings, diagnostic formulation, and treatment recommendations or plan. Must be performed by a physician, psychiatrist, LPHA, or master's-level CAP."),
    ],
)

# ===== 6. BPS/IDA 1-per-year =====
add_group(
    "6. BPS/IDA 'one per year' rule",
    "Appears in 3 places. §6 says it twice within 8 paragraphs. §FASD Billing Workflow tracks on the state fiscal year specifically.",
    [
        ("§6 Timelines",
         "A reminder is in Lauris for the following year: another BPS or IDA cannot be conducted until one year from the original date of completion."),
        ("§6 Insurance Requirements",
         "Another BPS or IDA cannot be conducted until one year from the original completion date."),
        ("§FASD Billing Workflow - Annual Reassessment",
         "Once per state fiscal year (July 1-June 30), complete one biopsychosocial evaluation OR one in-depth assessment. Track caps on the state fiscal year, not on the client's anniversary date of intake. Once an in-depth assessment is on file for a recipient, a biopsychosocial may not be billed afterward for that same recipient."),
    ],
)

# ===== 7. TP Review cadence (CONFLICT) =====
add_group(
    "7. Treatment Plan Review cadence - CONFLICT (90 days vs. every 6 months)",
    "§7 contradicts itself: one bullet says every 90 days, another says 'as frequently as every 3 months but not less than every 6 months.' The newer billing workflow says at-least-every-6-months with a 4-per-year cap.",
    [
        ("§7 Treatment Plan Deadline in Practice",
         "Treatment Plan Reviews must be completed every 90 days."),
        ("§7 Example Timeline",
         "Treatment Plan Reviews every 90 days."),
        ("§7 Required Content and Signatures",
         "Treatment Plans are reviewed as frequently as every 3 months but not less than every 6 months. Once the plan is done, put the review date on your calendar."),
        ("§FASD Billing Workflow - Treatment Plan Review",
         "A formal review is required at least every 6 months and may occur more often when significant changes occur. Reimbursement cap: 4 reviews per recipient per state fiscal year. A quarterly or every-4-months cadence (3-4 reviews per year) is compliant."),
    ],
)

# ===== 8. Pull up client in Lauris =====
add_group(
    "8. 'Pull up client in Lauris before billing'",
    "Appears in 2 places, including a floating bullet at the end of §4 that looks misplaced.",
    [
        ("§3.2 Service Codes intro",
         "Reference for the billing codes used at The Florida Center FASD Mental Health Division. Always pull up the client in Lauris and confirm the type of insurance and the service code available before billing."),
        ("§4 floating bullet at end",
         "Always pull up the client in Lauris and confirm the type of insurance and service code available before billing."),
    ],
)

# ===== 9. Complaints about therapy =====
add_group(
    "9. Complaints about therapy",
    "Appears in 2 places.",
    [
        ("§4 Complaints About Therapy",
         "Always encourage both the caregiver and the child to talk to you directly if something in therapy is not working. If they feel they cannot talk to you, they can contact the agency."),
        ("§5 First Session topics",
         "How to raise complaints about therapy. Always encourage the caregiver and child to talk to you directly, and let them know they can also go to the agency."),
    ],
)

# ===== 10. Prenatal Alcohol Screening recording =====
add_group(
    "10. Required training: Prenatal Alcohol Screening recording",
    "Appears in 2 places. Per the topic-ownership decision, Appendix B owns all training links.",
    [
        ("§6 Required Training on Prenatal Alcohol Screening",
         "Watch: The Importance of Screening for Prenatal Alcohol Exposure (Zoom recording) Passcode: SCREEN#02242025"),
        ("Appendix B Required Training Recordings",
         "The Importance of Screening for Prenatal Alcohol Exposure (Zoom recording) - Passcode: SCREEN#02242025"),
    ],
)

# ===== 11. Blank assessments / Box / Drive =====
add_group(
    "11. Blank assessments / Box / Drive resources",
    "Appears in 2 places. Per the topic-ownership decision, Appendix B owns all Box / Drive links.",
    [
        ("§6 Blank Forms and Example Materials",
         "Blank Assessments (Google Drive); Assessment and Diagnosis (Box); Observation (Box)"),
        ("Appendix B Assessment Resources",
         "Blank Assessments (Google Drive); Assessment and Diagnosis (Box); Observation (Box)"),
    ],
)

# ===== 12. Example DAP / BPS / TP / Discharge =====
add_group(
    "12. Example DAP Notes / BPS / TP / Discharge Plan link",
    "Appears in 2 places.",
    [
        ("§8 Example Documents",
         "Example DAP Notes, BPS, Treatment Plan, Discharge Plan"),
        ("Appendix B Example Clinical Documents",
         "Example DAP Notes, BPS, Treatment Plan, Discharge Plan (Google Drive); DAP Note Examples (Google Drive)"),
    ],
)

# ===== 13. Medicaid July 1 fiscal-year reset =====
add_group(
    "13. Medicaid July 1 fiscal-year reset / unit caps",
    "Appears in 2 places.",
    [
        ("§3.10 Unit Availability",
         "Medicaid units reset each July 1. Returning clients and transfers do not receive a new set of Medicaid units. Check the client's prior history for remaining units. Medicaid: 104 units (26 visits at 60 minutes) per fiscal year for H2019 HR. Does not include Assessment, Treatment Plan, or T1027 Psychoeducation."),
        ("§FASD Billing Workflow - Treatment Plan / Annual Reassessment",
         "Reimbursement cap: 1 treatment plan development per provider, with a maximum of 2 treatment plans per recipient, per state fiscal year (July 1-June 30). Reimbursement cap: 4 reviews per recipient per state fiscal year. Once per state fiscal year (July 1-June 30), complete one biopsychosocial evaluation OR one in-depth assessment."),
    ],
)

# ===== 14. H0032 Treatment Plan 1 per FY =====
add_group(
    "14. H0032 Treatment Plan 'one per fiscal year'",
    "Appears in 2 places.",
    [
        ("§3.2 H0032",
         "H0032 - Master Treatment Plan. 1 per recipient per fiscal year. Block 1 hour on the scheduler; billed per event."),
        ("§FASD Billing Workflow - Treatment Plan",
         "Reimbursement cap: 1 treatment plan development per provider, with a maximum of 2 treatment plans per recipient, per state fiscal year (July 1-June 30)."),
    ],
)

# ===== 15. Discharge / close form =====
add_group(
    "15. Discharge / close form",
    "Appears in 2 places. §3.11 is brief reminders; §10 is the full procedure. Probably keep §10 and delete §3.11, or convert §3.11 to a cross-reference.",
    [
        ("§3.11 Discharge Reminders",
         "When discharging a client, complete the close form request to remove the client from your caseload and discharge them in Lauris. If a discharged client still appears on your caseload, contact Durae. Please try to close all clients before the 1st of each month."),
        ("§10 Closing a Case",
         "(Full discharge workflow: Discharge Plan Requirements, Lauris Opening/Closing a Case, Lauris Discharge Plan, Discharge Workflow, Lauris Billing for Discharge - see §10 in full.)"),
    ],
)

doc.save(PATH)
print("Saved. Backup at FASD_Mental_Health_Onboarding_Manual_REVISED_BACKUP_pre_dup_review_2026-05-01.docx")
