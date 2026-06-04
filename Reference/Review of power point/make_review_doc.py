from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ---------- TITLE ----------
title = doc.add_heading('Review of "How to Bill" PowerPoint', level=0)

sub = doc.add_paragraph()
sub.add_run('Compared against the 2014 Florida Medicaid Community Behavioral Health Services Handbook and the current Florida Administrative Code (F.A.C.) coverage policies and fee schedule.').italic = True

doc.add_paragraph()

# ---------- SECTION: Framing ----------
doc.add_heading('First, an important framing issue', level=1)

p = doc.add_paragraph()
p.add_run('The 2014 Medicaid Handbook in your folder has been repealed as standalone authority. ').bold = True
p.add_run('The governing Florida Administrative Code rule that incorporated it (Rule 59G-4.050) is repealed. The current authority is split across several newer coverage policies, and rates/unit limits now live in the 2025 Community Behavioral Health Fee Schedule incorporated by reference in Rule 59G-4.002.')

# Authority table
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Rule'
hdr[1].text = 'Policy'
hdr[2].text = 'Effective'

auth_rows = [
    ('59G-4.028', 'Behavioral Health Assessment Services Coverage Policy', 'Nov 2019'),
    ('59G-4.052', 'Behavioral Health Therapy Services Coverage Policy', 'Nov 2019'),
    ('59G-4.031', 'Behavioral Health Community Support Services', 'current'),
    ('59G-4.370', 'Behavioral Health Intervention Services (TBOS, day services)', 'current'),
    ('59G-4.002', 'Provider Reimbursement Schedules & Billing Codes (2025 CBH Fee Schedule)', '2025'),
    ('59G-1.057', 'Telemedicine (defines what is/isn\'t reimbursable)', 'current'),
    ('65D-30.0044', 'Plans, Progress Notes, and Summaries (Substance Abuse) — treatment plan completion deadlines and QP countersignature', 'current'),
    ('65E-4.014', 'Community Mental Health — Standards for Client Records, Treatment and Quality Assurance', 'REPEALED 6-19-25'),
    ('59G-4.050', 'Community Behavioral Health Services (OLD 2014 handbook)', 'REPEALED'),
]
for r in auth_rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()
doc.add_paragraph(
    'Because F.A.C. trumps the old handbook, everything below is checked against BOTH the 2014 handbook and the current rules. Where they differ, the current rules govern.'
)

# ---------- SECTION: What the PPT gets right ----------
doc.add_heading('What the PowerPoint gets RIGHT', level=1)

right_items = [
    'Rule #1 — consent required before services (59G-4.028 §6 and 2014 handbook p. 2-3).',
    'Assessment must be billed first, before other services (handbook General Coverage §Assessment Requirement; 59G-4.028 §4.2).',
    'BBSE — H2010 HO, 15 or 30 min (1 or 2 units), max 2 units/day, max 10 units/year, cannot be same day as PSY or IDA. Confirmed by the 2025 fee schedule.',
    'Individual/Family therapy code H2019 HR — 1–4 units of 15/30/45/60 min; max 4 units/day. Confirmed by 2025 fee schedule.',
    'Group therapy code H2019 HQ.',
    'CFARS billed under H0031 (no modifier) for the first 3 per year; 4th not Medicaid-reimbursable and is shifted to CCM. Matches the handbook max of 3 Limited Functional Assessments per state fiscal year.',
    'ITXXCC for phone/collateral work, CCM for special reports (court, etc.) — consistent with the handbook\'s excluded services list (phone conversations and non-face-to-face activities aren\'t Medicaid-reimbursable).',
    'Family session without the client present is billable — 59G-4.052 §4.2.3 explicitly allows it: "when the recipient is not present, the services must always focus on the recipient."',
]
for item in right_items:
    doc.add_paragraph(item, style='List Bullet')

# ---------- SECTION: What the PPT gets WRONG ----------
doc.add_heading('What the PowerPoint gets WRONG or is misleading', level=1)

def add_issue(num, heading_text, body_paragraphs):
    h = doc.add_heading(f'{num}. {heading_text}', level=2)
    for para in body_paragraphs:
        if isinstance(para, list):
            # bullets
            for b in para:
                doc.add_paragraph(b, style='List Bullet')
        else:
            doc.add_paragraph(para)

add_issue(
    1,
    'PSY and IDA billed as "1 unit = 1 hour" — INCORRECT',
    [
        'Both the bio-psychosocial (H0031 HN) and the in-depth assessments (H0031 HO / H0031 TS) are billed per event / per assessment, not per hour. Per the 2025 fee schedule:',
        [
            'H0031 HN (PSY) = $57.28 per assessment (flat).',
            'H0031 HO (IDA, new patient) = $126.11 per assessment (flat).',
            'H0031 TS (IDA, established patient) = $100.88 per assessment (flat).',
        ],
        'Time spent does not change the reimbursement. The slide\'s "1 unit = 1 hour" framing implies time-based billing that does not exist for these codes.',
    ],
)

add_issue(
    2,
    '"ONE IDA per year AND ONE IDA Annual Update per year" — INCORRECT',
    [
        'The 2025 fee schedule and 59G-4.028 cap in-depth assessments (HO and TS combined) at ONE per recipient per state fiscal year. You cannot bill a new IDA (HO) and an established IDA (TS) in the same year.',
        'H0031 TS exists for recipients who need a new in-depth later (e.g., returning established patient) — it is NOT an "annual update" that can be stacked with HO.',
        'The same issue applies to the "PSY Annual Update — H0031 HN" wording. There is just one H0031 HN allowed per recipient per year, full stop.',
    ],
)

add_issue(
    3,
    'Treatment Plan / TPR "1 unit = 15, 30, 45, 60 min" — INCORRECT',
    [
        'H0032 (Master Treatment Plan) and H0032 TS (Treatment Plan Review) are both billed per event — one flat fee regardless of time:',
        [
            'H0032 = $97.86 per event.',
            'H0032 TS = $48.93 per event.',
        ],
        'The slide\'s time-tiered units are wrong. They bill as one event.',
    ],
)

add_issue(
    4,
    '"Medicaid allows 7 mins. for documentation, add this time to your session and round up" — WRONG and a compliance risk',
    [
        'The handbook\'s actual rule (p. 3-1, Units of Service) is a rounding rule applied to ACTUAL service time: if the last digit of total minutes is 7 or less, round down; if 8 or more, round up. Example from the handbook: 37 min = 2 units; 38 min = 3 units.',
        'There is no "7 free documentation minutes to add on." Billing for time not actually spent delivering the service is a billing-integrity problem and could constitute improper billing. This slide needs to be rewritten immediately.',
    ],
)

add_issue(
    5,
    '"1 additional TPR is allowed throughout the year if needed" — needs clarification',
    [
        'The handbook and 2025 fee schedule cap TPRs at a maximum of 4 per state fiscal year, period. The PowerPoint\'s timeline graphic shows 3 quarterly TPRs + 1 extra = 4, which matches. But the wording reads like "one per quarter (= 4) + 1 additional = 5." Tighten the language to: maximum 4 TPRs per state fiscal year.',
    ],
)

add_issue(
    6,
    'CFARS "1 unit = 15 minutes" — INCORRECT',
    [
        'H0031 Limited Functional Assessment is billed per assessment ($17.90 flat, 2025 fee schedule), not per 15-minute unit.',
        'Also, the state limit is 3 per year (not "1 per quarter" as a rule — though 3 quarterly does fit within the 3/year cap).',
    ],
)

add_issue(
    7,
    'BBSE used for phone crisis calls — compliance concern',
    [
        'The slide states: "Licensed clinicians will be able to document a BBSE — H2010 HO if a call is received to assist in the crisis intervention."',
        'Per Rule 59G-1.057, telephone conversations are explicitly excluded from reimbursable telemedicine. A BBSE must be face-to-face or via qualifying telemedicine (audio + video, real-time, two-way). A phone-only interaction should not be billed as H2010 HO.',
        'The same caveat applies to the later slide that says "bill a session — H2019 HR" for support-call interventions. Zoom/video is OK. Phone-only is not reimbursable.',
    ],
)

add_issue(
    8,
    '"Group and Individual cannot occur on the same day" — not a state rule',
    [
        'No such restriction appears in the 2014 handbook or in 59G-4.052 between H2019 HR and H2019 HQ.',
        'Same-day prohibitions exist for specific combinations (psychosocial rehab vs. behavioral health day services; medication management vs. brief individual/group medical psychotherapy), but not between H2019 HR and H2019 HQ.',
        'If your agency enforces "no same-day group + individual," label it clearly as an agency or managed-care-plan policy, not a Medicaid rule.',
    ],
)

add_issue(
    9,
    '"Treatment Plans need to be signed by the caregiver" — oversimplified',
    [
        '59G-4.028 §4.2.9 requires the signature of the recipient, OR the recipient\'s guardian if the recipient is under 18.',
        'For DCF / foster children, the caseworker may sign if the parent\'s signature cannot be obtained.',
        'There are also documented exceptions (emergency, DJJ custody, and recipients aged 13+ experiencing an emotional crisis per section 394.4784, F.S.).',
        'The current slide misses these nuances.',
    ],
)

add_issue(
    10,
    '"Treatment Plan within 10 business days" — correct direction, but cite the right F.A.C. rule and use the right unit',
    [
        'Correction to my earlier note: the 10-day figure is NOT "just agency policy." Florida Administrative Code 65D-30.0044 (Plans, Progress Notes, and Summaries) is the source the state auditors cite. The key verbatim language:',
        [
            '"If the treatment plan is completed by other than a qualified professional, the treatment plan shall be reviewed, countersigned, and dated by a qualified professional within 10 calendar days of completion." — Rule 65D-30.0044(1)(a), F.A.C.',
            '"For day or night treatment, the treatment plan shall be completed prior to or within 10 calendar days of placement." — Rule 65D-30.0044(1)(a)5., F.A.C.',
            '"For intensive outpatient treatment and outpatient treatment, the treatment plan shall be completed prior to or within 30 calendar days of placement." — Rule 65D-30.0044(1)(a)6., F.A.C.',
        ],
        'Important caveats:',
        [
            'Rule 65D-30 is the SUBSTANCE ABUSE chapter. The 10-calendar-day deadline applies directly to SA day/night treatment placements, and the 10-calendar-day countersignature rule applies whenever a non-qualified professional writes the plan.',
            'For general community outpatient mental health, the prior DCF rule (65E-4.014) required 30 days, and that rule was REPEALED effective June 19, 2025. There is no longer a DCF community-MH rule with an explicit completion deadline; the 65D-30.0044 rule is therefore the clearest F.A.C. deadline a SAMH-contracted provider is cited under.',
            'F.A.C. uses CALENDAR days, not BUSINESS days. 10 business days is roughly 14 calendar days — which is longer than F.A.C. allows. If the PowerPoint still says "10 business days," it is less strict than the F.A.C., which is why it triggered an audit finding. The slide should read "10 CALENDAR days" to match 65D-30.0044.',
            'The Florida Medicaid manual\'s "45 days" is NOT a completion deadline. It is a reimbursement look-back: Medicaid will reimburse for services provided within 45 days prior to the treating practitioner\'s signature (Rule 59G-4.028 §4.2.9; 2014 handbook p. 2-15). F.A.C. trumps the Medicaid manual, so 65D-30.0044 (10 calendar days) is the operative deadline for applicable service types, and the Medicaid 45-day figure does not override it.',
            'Whether 10 calendar days or 30 calendar days applies depends on the funding source, service type, and who is writing the plan (QP vs. non-QP). For SA day/night and any plan written by a non-QP = 10 calendar days. For outpatient SA = 30 calendar days. For outpatient MH reimbursed by Medicaid = no explicit F.A.C. completion deadline as of mid-2025, but the 10-calendar-day rule still applies if the plan is written by a non-QP under 65D-30.0044(1)(a).',
        ],
        'Bottom line for the slide: change "10 business days" to "10 CALENDAR days" and add the F.A.C. citation (Rule 65D-30.0044, F.A.C.). That matches what the auditors were citing and what the agency corrected to.',
    ],
)

# ---------- SECTION: Missing annual/daily limits ----------
doc.add_heading('Annual / daily LIMITS missing from the PowerPoint', level=1)
doc.add_paragraph('These limits are not mentioned in the slide deck but clinicians WILL hit them. State fiscal year runs July 1 – June 30.')

limits_table = doc.add_table(rows=1, cols=3)
limits_table.style = 'Light Grid Accent 1'
h = limits_table.rows[0].cells
h[0].text = 'Service'
h[1].text = 'Code'
h[2].text = 'Limit (state fiscal year)'

limits_rows = [
    ('Bio-psychosocial (PSY)', 'H0031 HN', '1 per year'),
    ('In-depth assessment (new + established combined)', 'H0031 HO / TS', '1 per year'),
    ('Limited functional (CFARS)', 'H0031', '3 per year'),
    ('BBSE', 'H2010 HO', '10 quarter-hour units (2.5 hrs) per year; max 2 units/day'),
    ('Individual / Family therapy', 'H2019 HR', '104 quarter-hour units (26 hrs) per year; max 4 units/day'),
    ('Group therapy', 'H2019 HQ', '156 quarter-hour units (39 hrs) per year; max 4 units/day'),
    ('Treatment plan development', 'H0032', '1 per provider per year; max 2 per recipient per year'),
    ('Treatment plan review', 'H0032 TS', '4 per year'),
]
for r in limits_rows:
    row = limits_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()

# ---------- SECTION: Sunshine Health codes ----------
doc.add_heading('Sunshine Health codes (T1023 HA, T1027, H2014)', level=1)
doc.add_paragraph(
    'These are plan-specific (Sunshine Health contract), not state AHCA-mandated codes. The PowerPoint correctly calls them out as Sunshine-only. Verify the unit limits and rates with the current Sunshine Health provider manual, because the state handbook and fee schedule do not govern them.'
)

# ---------- SECTION: Summary recommendations ----------
doc.add_heading('Summary recommendations', level=1)

recs = [
    'Add a cover slide stating which authority the deck is grounded in (Rule 59G-4.028, 59G-4.052, 2025 fee schedule) and note the 2014 handbook is historical reference only.',
    'Fix the unit/event confusion on PSY, IDA, MTP, TPR, and CFARS — these are per event, not time-based.',
    'Remove or rewrite the "7 minutes for documentation" slide. That is not how rounding works, and it creates compliance exposure.',
    'Clarify telemedicine rules: phone-only is never reimbursable; video/Zoom is (per 59G-1.057).',
    'Tighten the IDA "annual update" language. The combined cap is 1 in-depth assessment per year.',
    'Add the annual-limits table so clinicians know when they are about to bump a ceiling.',
    'Distinguish agency policy vs. state rule for the 10-business-day TP deadline and any same-day group+individual prohibitions.',
]
for r in recs:
    doc.add_paragraph(r, style='List Number')

# ---------- SECTION: Addendum on the 10-day rule ----------
doc.add_heading('Addendum: where the 10-day treatment-plan rule actually lives in F.A.C.', level=1)
doc.add_paragraph(
    'The state auditors were correct that the 10-day rule is in the Florida Administrative Code. I missed this on my first pass because I checked only the AHCA Medicaid rules (Chapter 59G). The deadline actually lives in the DCF Substance Abuse chapter: Rule 65D-30.0044, F.A.C. The relevant verbatim excerpts are below, pulled directly from the current rule text.'
)

doc.add_heading('Rule 65D-30.0044(1)(a), F.A.C. — Treatment Plan', level=2)
q = doc.add_paragraph()
q.add_run(
    '"If the treatment plan is completed by other than a qualified professional, the treatment plan shall be reviewed, countersigned, and dated by a qualified professional within 10 calendar days of completion."'
).italic = True

doc.add_heading('Rule 65D-30.0044(1)(a) 1.–6. — Completion deadlines by service level', level=2)
deadline_items = [
    ('Long-term outpatient methadone detox / methadone MAT', '30 calendar days of placement'),
    ('Intensive inpatient treatment', '3 calendar days of placement'),
    ('Residential treatment level 1', '7 calendar days of placement'),
    ('Residential levels 2, 3, 4; day/night treatment with community housing', '15 calendar days of placement'),
    ('Day or night treatment', '10 calendar days of placement'),
    ('Intensive outpatient treatment and outpatient treatment', '30 calendar days of placement'),
]
dl_table = doc.add_table(rows=1, cols=2)
dl_table.style = 'Light Grid Accent 1'
h = dl_table.rows[0].cells
h[0].text = 'Service level'
h[1].text = 'Deadline to complete treatment plan'
for svc, dl in deadline_items:
    row = dl_table.add_row().cells
    row[0].text = svc
    row[1].text = dl

doc.add_heading('Rule 65D-30.0044(1)(b) — Treatment Plan Reviews', level=2)
q = doc.add_paragraph()
q.add_run(
    '"Treatment plan reviews shall be completed with each individual and shall be signed and dated by the individual within 30 calendar days of the completion of the treatment plan."'
).italic = True
doc.add_paragraph('And for reviews written by a non-QP:')
q2 = doc.add_paragraph()
q2.add_run(
    '"For all components, if the treatment plan reviews are not completed by a qualified professional, the review shall be countersigned and dated by a qualified professional within five calendar days of the review."'
).italic = True

doc.add_heading('What this means for your slide deck', level=2)
meaning = [
    'The PowerPoint says "10 business days." F.A.C. says "10 CALENDAR days." 10 business days is roughly 14 calendar days — that is longer than the rule allows. This is almost certainly what the auditors flagged.',
    'Change the slide language to "10 CALENDAR days from the date of assessment/placement" and cite Rule 65D-30.0044, F.A.C.',
    'The Medicaid manual\'s "45 days" figure is NOT a completion deadline. It is a reimbursement look-back — Medicaid will reimburse services delivered in the 45 days before the treating practitioner signs. F.A.C. is the higher authority on deadlines, and the 45-day figure does not extend or override the 10-calendar-day rule.',
    'Whether the 10-day deadline actually applies in a given case depends on service type and who writes the plan. The QP countersignature 10-day rule applies across the board whenever a non-QP writes the plan. The 10-day placement deadline is specific to day/night treatment. Outpatient SA is 30 days. Outpatient MH (Medicaid-funded) has no explicit F.A.C. completion deadline as of mid-2025 (65E-4.014 was repealed June 2025), but the non-QP 10-day countersignature rule in 65D-30.0044 is still the safest operating assumption because auditors are citing it.',
    'Short version: if a non-licensed/non-QP clinician writes the plan, the QP countersignature must happen within 10 calendar days. Period.',
]
for m in meaning:
    doc.add_paragraph(m, style='List Bullet')

# ---------- SECTION: Sources ----------
doc.add_heading('Sources', level=1)
sources = [
    ('Rule 59G-4.050 — Community Behavioral Health Services (Repealed)', 'https://www.flrules.org/gateway/ruleno.asp?id=59G-4.050'),
    ('Rule 59G-4.028 — Behavioral Health Assessment Services Coverage Policy (Nov 2019)', 'https://ahca.myflorida.com/content/download/5937/file/59G-4.028.pdf'),
    ('Rule 59G-4.052 — Behavioral Health Therapy Services Coverage Policy (Nov 2019)', 'https://ahca.myflorida.com/content/download/5942/file/59G-4.052.pdf'),
    ('2025 Community Behavioral Health Fee Schedule', 'https://ahca.myflorida.com/content/download/26145/file/2025%20Community%20Behavoir%20Health%20Fee%20Schedule.pdf'),
    ('Rule 59G-1.057 — Telemedicine', 'https://ahca.myflorida.com/content/download/7011/file/59G_1057_TELEMEDICINE.pdf'),
    ('Rule 59G-4.002 — Provider Reimbursement Schedules and Billing Codes', 'https://ahca.myflorida.com/site/medicaid/rules/rule-59g-4.002-provider-reimbursement-schedules-and-billing-codes'),
    ('Rule 65D-30.0044 — Plans, Progress Notes, and Summaries (DCF Substance Abuse)', 'https://www.law.cornell.edu/regulations/florida/Fla-Admin-Code-r-65D-30-0044'),
    ('Chapter 65D-30, F.A.C. (full chapter, DCF)', 'https://myflfamilies.com/sites/default/files/2023-07/65D-30%20-%207.28.2023.pdf'),
]
for name, url in sources:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}\n')
    p.add_run(url).italic = True

doc.add_paragraph()
note = doc.add_paragraph()
note.add_run(
    'PDFs of 59G-4.028, 59G-4.052, and the 2025 Community Behavioral Health Fee Schedule have been downloaded into this "Review of power point" folder for reference.'
).italic = True

out = r'C:\Users\Tamra\Documents\Mental Health\Review of power point\How to Bill - PowerPoint Review.docx'
doc.save(out)
print('Wrote:', out)
