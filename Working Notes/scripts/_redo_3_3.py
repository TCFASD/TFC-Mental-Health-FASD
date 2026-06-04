"""
§3.3 Optimizing Billing cleanup:
1. Delete the @Claude question note at the top of §3.3.
2. In the IDA/BPS-same-day bullet, strip the tacked-on sentence
   'You can bill Assessment/TP on the same day, assessment 60 minutes,
   15 treatment plan, 15 therapy' (duplicates list above).
3. Move §7's 'Treatment Plan and individual therapy (ITX) can be billed
   on the same day...' bullet into §3.3 as its own clean bullet.
4. Delete that bullet from §7.
"""
from docx import Document
from copy import deepcopy

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)

# ---------- locate the targets ----------
claude_note_el = None
ida_bps_p = None  # the §3.3 bullet that will have its tail trimmed
s3_3_anchor = None  # where to insert the moved §7 bullet (after IDA/BPS bullet)
s7_tp_itx_el = None  # the §7 paragraph to move

in_section_3_3 = False
in_section_7 = False

for p in doc.paragraphs:
    t = p.text

    # Section boundaries
    if "3.3 Optimizing Billing" in t:
        in_section_3_3 = True
        in_section_7 = False
        continue
    if in_section_3_3 and (t.strip().startswith("3.10") or t.strip().startswith("3.6")
                           or t.strip().startswith("3.11") or t.strip().startswith("3.12")
                           or "Unit Availability" in t):
        in_section_3_3 = False
    if "Billing Rules for Treatment Plans" in t:
        in_section_7 = True
        continue
    if in_section_7 and ("Getting the Treatment Plan Signed" in t
                         or t.strip().startswith("7.")):
        in_section_7 = False

    # Targets in §3.3
    if in_section_3_3:
        if claude_note_el is None and "@Claude, is there any duplication" in t:
            claude_note_el = p._element
        if ida_bps_p is None and "You can bill an individual session and an IDA/BPS on the same day" in t:
            ida_bps_p = p

    # Target in §7
    if in_section_7 and s7_tp_itx_el is None and \
       "Treatment Plan and individual therapy (ITX) can be billed on the same day" in t:
        s7_tp_itx_el = p._element
        s7_tp_itx_text = t

if claude_note_el is None:
    raise SystemExit("Could not find the §3.3 @Claude note.")
if ida_bps_p is None:
    raise SystemExit("Could not find the §3.3 IDA/BPS bullet.")
if s7_tp_itx_el is None:
    raise SystemExit("Could not find the §7 TP+ITX same-day bullet.")

print("Located all targets.")

# ---------- 1. Delete the @Claude note ----------
claude_note_el.getparent().remove(claude_note_el)

# ---------- 2. Trim the tacked-on sentence from the IDA/BPS bullet ----------
# We rebuild the runs of this paragraph with the cleaned text.
old_text = ida_bps_p.text
tail = "You can bill Assessment/TP on the same day"
idx = old_text.find(tail)
if idx == -1:
    raise SystemExit("Tacked-on sentence not found in IDA/BPS bullet.")
new_text = old_text[:idx].rstrip()
# Make sure it ends with a period
if not new_text.endswith("."):
    new_text += "."

# Clear all runs and add a single replacement run
for run in list(ida_bps_p.runs):
    run._element.getparent().remove(run._element)
ida_bps_p.add_run(new_text)
print(f"Trimmed §3.3 IDA/BPS bullet. New ending: '...{new_text[-80:]}'")

# ---------- 3. Move §7 TP+ITX bullet into §3.3 ----------
# Insert it AFTER the (now-trimmed) IDA/BPS bullet.
# Detach from §7 first, then place after the IDA/BPS paragraph element.
s7_tp_itx_el.getparent().remove(s7_tp_itx_el)
ida_bps_p._element.addnext(s7_tp_itx_el)
print(f"Moved §7 bullet into §3.3: '{s7_tp_itx_text[:80]}...'")

doc.save(PATH)
print("§3.3 cleanup done.")
