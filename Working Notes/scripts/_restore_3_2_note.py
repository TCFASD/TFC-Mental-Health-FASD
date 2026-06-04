"""
Restore the NOTE paragraph that was incorrectly removed from §3.2:
'NOTE: Phone-only is not reimbursable...'
Insert it as a List Bullet directly after the new service-codes table.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)

note_text = ("NOTE:  Phone-only is not reimbursable. Per Rule 59G-1.057 "
             "(Telemedicine), telephone conversations are excluded from "
             "reimbursable telemedicine. Sessions must be face-to-face or "
             "via real-time, two-way video (Zoom). Phone calls go to ITXXCC "
             "(non-billable).")

# Find the §3.2 table (the one whose first cell is "Code")
target_table = None
for tbl in doc.element.body.iter():
    if tbl.tag.endswith("}tbl"):
        first_cell_text = ""
        for tc in tbl.iter():
            if tc.tag.endswith("}tc"):
                first_cell_text = "".join(
                    t.text or "" for t in tc.iter() if t.tag.endswith("}t")
                )
                break
        if first_cell_text.strip() == "Code":
            target_table = tbl
            break

if target_table is None:
    raise SystemExit("Could not find the §3.2 service codes table.")

# Build a new List Bullet paragraph with the NOTE
note_p = doc.add_paragraph(style="List Bullet")
note_p.add_run(note_text)

# Move the new paragraph element to right after the table
target_table.addnext(note_p._element)

doc.save(PATH)
print("Restored §3.2 NOTE bullet directly under the service codes table.")
