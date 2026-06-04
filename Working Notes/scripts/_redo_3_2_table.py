"""
§3.2 Service Codes: convert the bulleted code list into a 3-column table
(Code | Service | Limits / Notes), and delete the @Claude meta-note.

Wording preserved as closely as possible. The 'phone-only is not reimbursable'
NOTE bullet stays in place and ends up directly under the table.

CFARS row keeps current §3.2 wording; left for the duplicate-review pass.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)

# 1. Find the intro paragraph (anchor for table insertion) and the
#    paragraphs to remove (@Claude note + 15 code rows).
intro_el = None
to_remove = []

removal_markers = [
    "@Claude, make this into a table",
    "H0031 HN ",
    "H0031 HO ",
    "H0031 TS ",
    "H0031 (no modifier)",
    "H0032 ",
    "H0032 TS ",
    "H2010 HO ",
    "H2019 HR ",
    "H2019 HQ ",
    "T1023 HA ",
    "T1027 ",
    "H2014 ",
    "90837 (commercial insurance)",
    "ITXXCC ",
    "CCM ",
]

# We need to scope the removal to §3.2 only (some of these tokens could
# theoretically appear elsewhere). Walk the body and only flag matches
# that occur between the §3.2 heading and the §3.4 (or §3.3) heading.
in_section = False
for p in doc.paragraphs:
    t = p.text
    if "3.2 Service Codes" in t:
        in_section = True
        continue
    if in_section and (t.strip().startswith("3.3") or t.strip().startswith("3.4")
                       or "3.4 Z Codes" in t or "3.3 Optimizing Billing" in t):
        in_section = False
        break
    if not in_section:
        continue
    if intro_el is None and "Always pull up the client in Lauris" in t:
        intro_el = p._element
        continue
    for m in removal_markers:
        if m in t:
            to_remove.append(p._element)
            break

if intro_el is None:
    raise SystemExit("Could not find §3.2 intro paragraph.")
print(f"Will remove {len(to_remove)} paragraphs from §3.2 "
      f"(expected 16: 1 @Claude note + 15 code rows).")

# 2. Build the table at end of doc, then move it after the intro paragraph.
rows = [
    ("H0031 HN", "PSY (Biopsychosocial Assessment)",
     "1 per recipient per fiscal year. Block 1 hour on the scheduler; billed per event."),
    ("H0031 HO", "IDA, new patient",
     "1 per recipient per fiscal year (HO + TS combined). Block 1 hour on the scheduler; billed per event."),
    ("H0031 TS", "IDA, established (returning) patient",
     "Counts toward the 1-per-year IDA cap with HO. 1 hour."),
    ("H0031 (no modifier)", "CFARS (Limited Functional Assessment)",
     "Medicaid only. The first 3 per state fiscal year are billed here; the 4th is billed to CCM. Client must be present."),
    ("H0032", "Master Treatment Plan",
     "1 per recipient per fiscal year. Block 1 hour on the scheduler; billed per event."),
    ("H0032 TS", "Treatment Plan Review",
     "Up to 4 per state fiscal year (3 quarterly + 1 additional if needed). Block 1 hour on the scheduler; billed per event."),
    ("H2010 HO", "BBSE (Brief Behavioral Status Exam)",
     "Licensed clinicians only. 1–2 units (15 or 30 minutes). Maximum 2 units per day, maximum 10 units per year. Cannot be billed the same day as PSY or IDA."),
    ("H2019 HR", "Individual / Family Therapy",
     "Billable in 15/30/45/60-minute units (1–4 units). Maximum 4 units per day. Cannot bill two H2019 HR codes in one day, regardless of length or who was present. Sessions can be billed as: individual (therapist with client alone), family with client present, or family without client present (focus must still be on the client's therapeutic work)."),
    ("H2019 HQ", "Group Therapy",
     "Workflow and same-day-as-individual rule pending – see questions.docx."),
    ("T1023 HA", "Sunshine Health, ages 0–5",
     "40 units (10 hours) per calendar year. Used for tests, inventories, questionnaires, and structured observations to assess the caregiver-child relationship and inform treatment plan development. Not used often because reimbursement is low; clinicians often prefer to incorporate this work into regular billing."),
    ("T1027", "Sunshine Health, ages 0–21 with SED diagnosis (parent coaching)",
     "Minimum bachelor's degree (master's-level clinicians can also perform it). The DAP note must NOT use the word \"therapy.\" If FMF is delivered for the first 60 minutes and T1027 is billed for the time after, the focus of the T1027 portion must be different – parent coaching, not therapy."),
    ("H2014", "Sunshine Health, Early Childhood Court (ECC), CPP, ages 0–3",
     "No limit. Used for CPP sessions, time in court, and time in family team meetings. Pilot started June 1, 2025."),
    ("90837", "Commercial insurance individual therapy",
     "Billed for 60-minute therapy sessions."),
    ("ITXXCC", "Clinical Consultation",
     "Used for phone calls with clients or collateral conversations with other support people (e.g., speech therapist, school counselor). Not billable to insurance."),
    ("CCM", "Clinical Case Management",
     "Used for special reports outside normal documentation (court reports, etc.). Not billable to insurance."),
]

table = doc.add_table(rows=1 + len(rows), cols=3)
table.style = "Table Grid"

# Header row
hdr = table.rows[0]
for cell, label in zip(hdr.cells, ("Code", "Service", "Limits / Notes")):
    # bold label
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    # light-blue header shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DEEAF6")  # light blue
    tcPr.append(shd)

# Body rows
for i, (code, service, notes) in enumerate(rows, start=1):
    cells = table.rows[i].cells
    # Code cell - bold
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(code)
    r0.bold = True
    r0.font.size = Pt(10)
    # Service
    p1 = cells[1].paragraphs[0]
    p1.add_run(service).font.size = Pt(10)
    # Notes
    p2 = cells[2].paragraphs[0]
    p2.add_run(notes).font.size = Pt(10)

# Move the table to be the next sibling of the intro paragraph
intro_el.addnext(table._element)

# 3. Remove @Claude note + 15 code paragraphs
removed = 0
for el in to_remove:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)
        removed += 1

doc.save(PATH)
print(f"§3.2 done: built {len(rows)}-row table; removed {removed} old paragraphs.")
