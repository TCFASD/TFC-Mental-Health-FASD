"""
Apply Section 2 + Section 3 + Section 7 reorganization to the FASD onboarding manual.
Backup already created. Operates on FASD_Mental_Health_Onboarding_Manual.docx in place.
"""
from docx import Document

source = 'FASD_Mental_Health_Onboarding_Manual.docx'
target = 'FASD_Mental_Health_Onboarding_Manual_REVISED.docx'
doc = Document(source)


def find_h1(prefix):
    """Find the index of the Heading 1 paragraph that starts with `prefix`."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and p.text.strip().startswith(prefix):
            return i
    return -1


sec2 = find_h1('2. ')
sec3 = find_h1('3. ')
sec4 = find_h1('4. ')
sec7 = find_h1('7. ')
sec8 = find_h1('8. ')

print(f'Section indices: 2={sec2}, 3={sec3}, 4={sec4}, 7={sec7}, 8={sec8}')
assert all(i > 0 for i in (sec2, sec3, sec4, sec7, sec8)), 'Heading detection failed'


def set_paragraph_text(p, new_text):
    """Replace the text of a paragraph while preserving its style/formatting."""
    if p.runs:
        for r in p.runs[1:]:
            r.text = ''
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


# === STAGE 1: Update §7 wording (10 business days -> 10 calendar days) ===
sec7_paras = doc.paragraphs[sec7:sec8]
for p in sec7_paras:
    if 'business' in p.text.lower():
        for run in p.runs:
            run.text = (run.text
                        .replace('10-business-day', '10-calendar-day')
                        .replace('10 business day', '10 calendar day')
                        .replace('10 business-day', '10 calendar-day')
                        .replace('10-business day', '10-calendar day'))
        if 'business' in p.text.lower():
            print(f'  WARN: §7 paragraph still contains "business": {p.text[:100]}')
        else:
            print(f'  §7 fixed: {p.text[:100]}')

# === STAGE 2: Capture anchors for §2 and §3 BEFORE deleting content ===
sec2_heading = doc.paragraphs[sec2]
sec3_anchor = doc.paragraphs[sec3]   # we'll insert §2 content before this
sec3_heading = doc.paragraphs[sec3]
sec4_anchor = doc.paragraphs[sec4]   # we'll insert §3 content before this

# Capture the paragraphs to delete
paras_delete_sec2 = list(doc.paragraphs[sec2 + 1:sec3])
paras_delete_sec3 = list(doc.paragraphs[sec3 + 1:sec4])

# === STAGE 3: Rename §2 heading ===
set_paragraph_text(sec2_heading, '2. FASD Families May Need More')

# §3 heading "3. Billing" stays as-is.

# === STAGE 4: Delete old §2 and §3 content ===
for p in paras_delete_sec2:
    remove_paragraph(p)
for p in paras_delete_sec3:
    remove_paragraph(p)

# === STAGE 5: Insert new §2 content (5 bare bullets) ===
def insert_before(anchor, text, style):
    return anchor.insert_paragraph_before(text, style=style)


for bullet in [
    'Educational Navigation',
    'APD',
    'Connecting with LRT',
    'CAT team',
    'Parent Advocate',
]:
    insert_before(sec3_anchor, bullet, 'List Bullet')


# === STAGE 6: Insert new §3 content ===
def add_h2(text):
    return insert_before(sec4_anchor, text, 'Heading 2')


def add_normal(text):
    return insert_before(sec4_anchor, text, 'Normal')


def add_bullet(text):
    return insert_before(sec4_anchor, text, 'List Bullet')


def add_bold_lead(bold_text, rest):
    """Paragraph that starts with a bold lead-in (e.g. a code name) followed by description."""
    p = insert_before(sec4_anchor, '', 'Normal')
    # Clear default run added by insert_paragraph_before
    if p.runs:
        for r in p.runs:
            r.text = ''
    r1 = p.add_run(bold_text + ' ')
    r1.bold = True
    p.add_run(rest)
    return p


# 3.1 Billing Absolutes
add_h2('3.1 The Billing Absolutes — No Ifs, Ands, or Buts')
add_bullet('Rule #1. No client can be seen until consents are signed.')
add_bullet('Rule #2. The assessment is always billed first. Nothing else can be billed before it. (ITXXCC and CCM are exceptions because they are not billable codes — see the Service Codes in §3.2.)')
add_bullet('Rule #3. You must complete a Treatment Plan within 10 calendar days from the date of the assessment, and the caregiver must sign it. (Per Rule 65D-30.0044, F.A.C. Note: this is calendar days, not business days.)')
add_bullet('Rule #4. You must track when documents expire and cannot see clients until they have been renewed. This applies to consents, assessments, treatment plans, and treatment plan reviews.')

# 3.2 Service Codes
add_h2('3.2 Service Codes')
add_normal('Reference for the billing codes used at The Florida Center FASD Mental Health Division. Always pull up the client in Lauris and confirm the type of insurance and the service code available before billing.')

codes = [
    ('H0031 HN — PSY (Biopsychosocial Assessment).',
     '1 per recipient per fiscal year. Block 1 hour on the scheduler; billed per event.'),
    ('H0031 HO — IDA, new patient.',
     '1 per recipient per fiscal year (HO + TS combined). Block 1 hour on the scheduler; billed per event.'),
    ('H0031 TS — IDA, established (returning) patient.',
     'Counts toward the 1-per-year IDA cap with HO. 1 hour.'),
    ('H0031 (no modifier) — CFARS (Limited Functional Assessment).',
     'Medicaid only. The first 3 per state fiscal year are billed here; the 4th is billed to CCM. Client must be present.'),
    ('H0032 — Master Treatment Plan.',
     '1 per recipient per fiscal year. Block 1 hour on the scheduler; billed per event.'),
    ('H0032 TS — Treatment Plan Review.',
     'Up to 4 per state fiscal year (3 quarterly + 1 additional if needed). Block 1 hour on the scheduler; billed per event.'),
    ('H2010 HO — BBSE (Brief Behavioral Status Exam).',
     'Licensed clinicians only. 1–2 units (15 or 30 minutes). Maximum 2 units per day, maximum 10 units per year. Cannot be billed the same day as PSY or IDA.'),
    ('H2019 HR — Individual / Family Therapy.',
     'Billable in 15/30/45/60-minute units (1–4 units). Maximum 4 units per day. Cannot bill two H2019 HR codes in one day, regardless of length or who was present. Sessions can be billed as: individual session (therapist with client alone), family session with client present, or family session without client present (the focus must still be on the client’s therapeutic work).'),
    ('H2019 HQ — Group Therapy.',
     'Workflow and same-day-as-individual rule pending — see questions.docx.'),
    ('T1023 HA — Sunshine Health, ages 0–5.',
     '40 units (10 hours) per calendar year. Used for tests, inventories, questionnaires, and structured observations to assess the caregiver-child relationship and inform treatment plan development. Not used often because the reimbursement is low; clinicians often prefer to incorporate this work into regular billing.'),
    ('T1027 — Sunshine Health, ages 0–21 with SED diagnosis.',
     'Minimum bachelor’s degree (master’s-level clinicians can also perform it). Used for parent coaching to help caregivers understand and manage the child’s behavioral needs. The DAP note must NOT use the word “therapy.” If FMF is delivered for the first 60 minutes and T1027 is billed for the time after, the focus of the T1027 portion must be different — parent coaching, not therapy. Submit an authorization request to Durae after an assessment with a documented diagnosis (CBHA, IDA, BPS, or TP). Unit cap pending verification — see questions.docx.'),
    ('H2014 — Sunshine Health, Early Childhood Court (ECC), CPP, ages 0–3.',
     'No limit. Used for CPP sessions, time in court, and time in family team meetings. Pilot started June 1, 2025.'),
    ('ITXXCC — Clinical Consultation.',
     'Used for phone calls with clients or collateral conversations with other support people (e.g., speech therapist, school counselor). Not billable to insurance.'),
    ('CCM — Clinical Case Management.',
     'Used for special reports outside normal documentation (court reports, etc.). Not billable to insurance.'),
]

for code_title, code_desc in codes:
    add_bold_lead(code_title, code_desc)

# 3.3 First Encounter
add_h2('3.3 What You Can Bill on First Encounter')
add_normal('On the day of the PSY (or IDA), a clinician can bill all three of the following:')
add_bullet('ONE PSY or IDA — H0031 HN — 1 unit')
add_bullet('ONE TP — H0032 — 1 unit')
add_bullet('ONE CFARS — H0031 — 1 unit (Medicaid only; client must be present)')
add_normal('Ask your supervisor about how to optimize this time.')

# 3.4 Z Codes
add_h2('3.4 Z Codes')
add_bullet('Z codes are not billable codes and cannot serve as the primary diagnosis if the client has insurance coverage.')
add_bullet('If a Z code is the sole diagnosis, consult with your supervisor before submitting any notes.')

# 3.5 Billing Restrictions
add_h2('3.5 Billing Restrictions')
add_bullet('90837 (commercial insurance) is billed for 60-minute therapy sessions.')
add_bullet('H2019 HR (Medicaid) is billable in 15/30/45/60-minute units. Maximum 4 units (60 minutes) per day. Cannot bill two H2019 HR codes in one day, regardless of length or who was present.')
add_bullet('Therapy unit rounding: round based on actual service time. If the last digit of total minutes is 7 or less, round down; if 8 or more, round up. Example: 37 minutes = 2 units; 38 minutes = 3 units.')
add_bullet('Phone-only is not reimbursable. Per Rule 59G-1.057 (Telemedicine), telephone conversations are excluded from reimbursable telemedicine. Sessions must be face-to-face or via real-time, two-way video (Zoom). Phone calls go to ITXXCC (non-billable).')
add_bullet('BBSE (H2010 HO) limits — see Service Codes in §3.2.')
add_bullet('Always pull up the client in Lauris and confirm the type of insurance and service code available before billing.')
add_bullet('Bob runs reports on visit counts.')

# 3.6 Insurance Verification
add_h2('3.6 Insurance Verification')
add_normal('Clients in the Insurance Verification caseload will have units added by Durae. Notify her if immediate units are needed; otherwise allow 48 hours once a client has been entered in Lauris.')

# 3.7 Client Transfers and Dual Enrollment
add_h2('3.7 Client Transfers and Dual Enrollment')
add_bullet('Clients can only be transferred within the same program to a new therapist and/or school.')
add_bullet('Clients cannot be transferred between programs. Submit a close form and inform Intake and Durae to open them in the new program (e.g., SBMH to OPMH).')
add_bullet('Clients cannot be enrolled in two mental health programs simultaneously (e.g., ECC and OPMH). Exceptions require supervisor approval case-by-case.')

# 3.8 Document Uploads and Intake Forms
add_h2('3.8 Document Uploads and Intake Forms')
add_bullet('Scan and email Intake with the client’s name and ID number for any required document uploads.')
add_bullet('Send documents separately (ACE, Consent Forms, Referral, etc.).')
add_bullet('Intake forms must be legible with correct name spelling, DOB, phone, open date, and reason for referral.')
add_bullet('The insurance portion must not be left blank. If you lack the info, use that space to tell Intake what you need (e.g., “Please look up client in portal” or “Please contact parent”).')
add_bullet('If no PCP name is provided, do not send the form to Intake.')
add_bullet('Consent forms must be obtained prior to or on the day of any scheduled BPS, IDA, or TP appointment. (See §4 for full consent requirements.)')

# 3.9 What Intake Does Before Reaching the Therapist
add_h2('3.9 What Intake Does Before Reaching the Therapist')
add_bullet('Consents, pediatric intake, medical document.')

# 3.10 Unit Availability, Updates, and Service Limits
add_h2('3.10 Unit Availability, Updates, and Service Limits')
add_bullet('For unit availability issues, contact Durae with the client ID, date of service, and code in question.')
add_bullet('Notify Intake and Durae of any client updates (name, address, phone, insurance).')
add_bullet('If a client resides in Sarasota County, one of the billing codes must be G-FETAL ALCOHOL — COUNTY, and the FASD consult/case management must be under that billing code. This is NOT applicable outside Sarasota County.')
add_bullet('Medicaid units reset each July 1. Returning clients and transfers do not receive a new set of Medicaid units. Check the client’s prior history for remaining units.')
add_bullet('Medicaid: 104 units (26 visits at 60 minutes) per fiscal year for H2019 HR. Does not include Assessment, Treatment Plan, or T1027 Psychoeducation.')

# 3.11 Discharge Reminders
add_h2('3.11 Discharge Reminders')
add_bullet('When discharging a client, complete the close form request to remove the client from your caseload and discharge them in Lauris.')
add_bullet('If a discharged client still appears on your caseload, contact Durae.')
add_bullet('Please try to close all clients before the 1st of each month.')

# 3.12 Payor Usage
add_h2('3.12 Payor Usage (G-Fetal Alcohol County vs. Srq Office)')
add_normal('The following rules clarify which payor to use when billing FASD-related services. Confirm with Charmian if a situation is not covered here.')
add_normal('If Jackie is running out of units for continued service and we are unable to get approval for additional units, “G-Fetal Alcohol-Srq Office” can be used to support the work being done.')
add_normal('Regarding the use of “FASD Unassigned Client” with “G-Fetal Alcohol County”: the safest approach is to discontinue this practice. If you need to track work being done with individuals who are not yet clients, you may continue using “FASD Unassigned Client,” but the payor must be “G-Fetal Alcohol-Srq Office.”')
add_normal('Quick reference:')
add_bullet('“G-Fetal Alcohol County” — only to be used by Jackie for intervention with an established Sarasota County client. May also be used by the evaluation team when completing an evaluation for a confirmed Sarasota County resident.')
add_bullet('“G-Fetal Alcohol-Srq Office” — use for all other situations not noted above. This is the only payor to be used with “FASD Unassigned Client.”')

# === STAGE 7: Update TOC entries that we've changed ===
toc_updates = {
    '2. FASD and Neurobehavioral Conditions: An Overview': '2. FASD Families May Need More',
}

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in toc_updates:
        set_paragraph_text(p, toc_updates[txt])
        print(f'  TOC updated: {txt!r} -> {toc_updates[txt]!r}')

# === STAGE 8: Save to a NEW revised file (original untouched) ===
doc.save(target)
print(f'\nSaved revised copy to: {target}')
print(f'Original {source} is unchanged.')
