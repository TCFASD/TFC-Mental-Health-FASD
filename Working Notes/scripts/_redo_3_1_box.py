"""
§3.1 Billing Absolutes: replace the bulleted rules with a decorative
shaded/bordered single-cell box, and delete the @Claude meta-note.

Rule wording is preserved verbatim. Rule #3 (10 calendar days) is left
untouched - it is part of the duplicate-review conflict group #4.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"C:\Users\Tamra\Documents\Mental Health\FASD_Mental_Health_Onboarding_Manual_REVISED.docx"

doc = Document(PATH)

# 1. Locate the heading and the paragraphs to remove (by their text)
heading_el = None
to_remove = []  # list of paragraph elements

markers = {
    "claude_note": "@Claude, make this more decorative",
    "rule1": "Rule #1. No client can be seen until consents are signed",
    "rule2": "Rule #2. The assessment is always billed first",
    "rule3": "Rule #3. You must complete a Treatment Plan within 10 calendar days",
    "rule4": "Rule #4. You must track when documents expire",
}
found = {}
for p in doc.paragraphs:
    if "3.1 The Billing Absolutes" in p.text and heading_el is None:
        heading_el = p._element
        continue
    for key, m in markers.items():
        if m in p.text and key not in found:
            found[key] = p._element

if heading_el is None:
    raise SystemExit("Could not find §3.1 heading - aborting.")
missing = [k for k in markers if k not in found]
if missing:
    raise SystemExit(f"Could not find: {missing}")

to_remove = list(found.values())

# 2. Build a 1x1 bordered, shaded table at the end of the doc
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)

# Cell shading: light yellow
tcPr = cell._tc.get_or_add_tcPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"), "FFF2CC")
tcPr.append(shd)

# Cell borders: thick black on all sides
tcBorders = OxmlElement("w:tcBorders")
for side in ("top", "left", "bottom", "right"):
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "12")  # 1.5pt-ish
    b.set(qn("w:color"), "000000")
    tcBorders.append(b)
tcPr.append(tcBorders)

# Cell content: bold header + 4 rules
# Reuse the existing default paragraph for the header
hdr_p = cell.paragraphs[0]
hdr_run = hdr_p.add_run("BILLING ABSOLUTES — NO IFS, ANDS, OR BUTS")
hdr_run.bold = True
hdr_run.font.size = Pt(13)

rules_text = [
    "Rule #1.  No client can be seen until consents are signed.",
    "Rule #2.  The assessment is always billed first. Nothing else can be billed before it. "
    "(ITXXCC and CCM are exceptions because they are not billable codes – see the Service Codes in §3.2.)",
    "Rule #3.  You must complete a Treatment Plan within 10 calendar days from the date of the "
    "assessment, and the caregiver must sign it. (Per Rule 65D-30.0044, F.A.C. Note: this is "
    "calendar days, not business days.)",
    "Rule #4.  You must track when documents expire and cannot see clients until they have been "
    "renewed. This applies to consents, assessments, treatment plans, and treatment plan reviews.",
]
for rt in rules_text:
    p = cell.add_paragraph()
    r = p.add_run(rt)
    r.bold = False
    r.font.size = Pt(11)

# 3. Move the table to be the next sibling of the §3.1 heading
heading_el.addnext(table._element)

# 4. Remove the @Claude note + 4 rule paragraphs from their original spots
for el in to_remove:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

doc.save(PATH)
print("§3.1 done: rules now in a shaded bordered box; @Claude note deleted.")
